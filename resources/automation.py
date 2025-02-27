import os
import pandas as pd
import docx
import pypdf

from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from docling.document_converter import DocumentConverter

# Carregar variáveis do .env
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

# Criar modelo LLM para resumo
llm = ChatOpenAI(model_name="gpt-4", openai_api_key=openai_api_key)

# Criar conversor de documentos com DocLing
converter = DocumentConverter()

# Função para extrair dados de documentos
def extract_data(file_path):
    """Processa um documento e retorna seu conteúdo extraído."""
    if file_path.endswith((".pdf", ".docx", ".xlsx", ".csv", ".txt")):
        try:
            print(f"📂 Processando: {file_path}")

            result = converter.convert(file_path)
            if result.status.value.lower() != "success":
                print(f"⚠ Erro ao processar {file_path}: {result.status}")
                return None

            doc = result.document  # Pegamos o documento processado pelo Docling

            extracted_text = []

            # Percorre os textos do corpo do documento
            for text_item in doc.texts:
                if hasattr(text_item, 'text') and text_item.text:  # Verifica se tem conteúdo
                    extracted_text.append(text_item.text.strip())

            print(f"🔍 Resultado da conversão: {extracted_text}")  # Debug

            # Junta todo o texto extraído com quebras de linha
            return "\n".join(extracted_text)

        except Exception as e:
            print(f"⚠ Erro ao processar {file_path}: {e}")
            return ""
    return ""

# Função para gerar um relatório estruturado
def generate_report(data):
    """Gera um relatório formatado a partir dos dados extraídos."""
    doc = docx.Document()
    doc.add_heading("Relatório Automatizado", level=1)
    doc.add_paragraph(data)
    
    report_path = "output/relatorio_automatizado.docx"
    doc.save(report_path)
    print(f"📄 Relatório salvo em: {report_path}")

# Função para resumir conteúdo usando IA
def summarize_text(text):

    MAX_TOKENS = 4000  # Definir um limite para cada requisição ao GPT
    chunks = [text[i:i+MAX_TOKENS] for i in range(0, len(text), MAX_TOKENS)]

    """Resume um texto longo utilizando IA."""
    try:
        print("⏳ Gerando resumo...")
        summaries = []
        for chunk in chunks:
            response = llm.invoke(chunk)  # Usando `invoke()` em vez de `__call__`
            summaries.append(response.content)

        return " ".join(summaries)
    except Exception as e:
        print(f"⚠ Erro ao resumir texto: {e}")
        return "Erro ao gerar resumo."

# Criar pasta de saída se não existir
if not os.path.exists("output"):
    os.makedirs("output")

# Testar o fluxo
if __name__ == "__main__":
    # Documento para teste
    test_file = "data/exemplo.docx"

    # Extração
    text_data = extract_data(test_file)
    if text_data:
        print("✅ Dados extraídos com sucesso!")

        # Resumo do conteúdo
        summary = summarize_text(text_data)
        print("📌 Resumo Gerado:\n", summary)

        # Geração de relatório
        generate_report(summary)
